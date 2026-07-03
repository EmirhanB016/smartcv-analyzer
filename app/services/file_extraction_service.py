"""Extract readable text from supported CV file formats."""

from io import BytesIO

import fitz
from docx import Document
from docx.document import Document as DocxDocument

from app.schemas.errors import ErrorCode, raise_api_error
from app.utils.validators import ValidatedCVFile


def extract_cv_text(cv_file: ValidatedCVFile) -> str:
    """Extract text from a validated CV file."""
    if cv_file.extension == ".pdf":
        return extract_pdf_text(cv_file.content)

    if cv_file.extension == ".docx":
        return extract_docx_text(cv_file.content)

    raise_api_error(ErrorCode.UNSUPPORTED_FILE_TYPE)


def extract_pdf_text(content: bytes) -> str:
    """Extract text from PDF bytes with PyMuPDF."""
    document = None
    try:
        document = fitz.open(stream=content, filetype="pdf")
        text_parts = [page.get_text("text") for page in document]
    except Exception as exc:
        raise_api_error(ErrorCode.TEXT_EXTRACTION_FAILED) from exc
    finally:
        if document is not None:
            document.close()

    return _require_readable_text("\n".join(text_parts))


def extract_docx_text(content: bytes) -> str:
    """Extract paragraph and table text from DOCX bytes."""
    try:
        document = Document(BytesIO(content))
        text_parts = _extract_docx_paragraphs(document)
        text_parts.extend(_extract_docx_tables(document))
    except Exception as exc:
        raise_api_error(ErrorCode.TEXT_EXTRACTION_FAILED) from exc

    return _require_readable_text("\n".join(text_parts))


def _extract_docx_paragraphs(document: DocxDocument) -> list[str]:
    return [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]


def _extract_docx_tables(document: DocxDocument) -> list[str]:
    table_lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append("\t".join(cells))

    return table_lines


def _require_readable_text(text: str) -> str:
    stripped_text = text.strip()
    if not stripped_text or not any(character.isalnum() for character in stripped_text):
        raise_api_error(ErrorCode.TEXT_EXTRACTION_FAILED)

    return stripped_text
