from io import BytesIO

import fitz
from docx import Document
from fastapi import HTTPException
import pytest

from app.services.file_extraction_service import (
    extract_cv_text,
    extract_docx_text,
    extract_pdf_text,
)
from app.utils.validators import ValidatedCVFile


def test_extract_docx_text_reads_paragraphs_and_table_cells() -> None:
    document = Document()
    document.add_paragraph("Jane Developer")
    document.add_paragraph("Python FastAPI Docker")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Education"
    table.cell(0, 1).text = "Computer Engineering"

    text = extract_docx_text(_docx_bytes(document))

    assert "Jane Developer" in text
    assert "Python FastAPI Docker" in text
    assert "Education" in text
    assert "Computer Engineering" in text


def test_extract_pdf_text_reads_page_text() -> None:
    text = extract_pdf_text(_pdf_bytes("Python FastAPI REST API developer"))

    assert "Python FastAPI REST API developer" in text


def test_extract_cv_text_dispatches_by_validated_extension() -> None:
    cv_file = ValidatedCVFile(
        filename="cv.docx",
        extension=".docx",
        content=_docx_bytes_with_paragraph("SmartCV Analyzer project"),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert "SmartCV Analyzer project" in extract_cv_text(cv_file)


def test_extract_docx_text_rejects_corrupted_content() -> None:
    with pytest.raises(HTTPException) as exc_info:
        extract_docx_text(b"not a real docx")

    assert exc_info.value.status_code == 422
    assert exc_info.value.detail["code"] == "TEXT_EXTRACTION_FAILED"


def test_extract_pdf_text_rejects_corrupted_content() -> None:
    with pytest.raises(HTTPException) as exc_info:
        extract_pdf_text(b"not a real pdf")

    assert exc_info.value.status_code == 422
    assert exc_info.value.detail["code"] == "TEXT_EXTRACTION_FAILED"


def _docx_bytes(document: Document) -> bytes:
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _docx_bytes_with_paragraph(text: str) -> bytes:
    document = Document()
    document.add_paragraph(text)
    return _docx_bytes(document)


def _pdf_bytes(text: str) -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text, fontsize=11)
    content = document.tobytes()
    document.close()
    return content
